"""
md_to_docx.py — convert a campaign markdown report (with base64-embedded or path PNGs) to .docx
via python-docx (pandoc is not installed). Handles the constructs the reports use: ATX headings
(#/##/###), paragraphs with **bold**/`code`, pipe tables, images ![alt](data:..|path), blockquotes
(>), and -/* bullet lists. Not a general Markdown engine — just enough for these reports.

Usage:  python -m validation.md_to_docx <report.md> [out.docx] [figs_base_dir]
"""
from __future__ import annotations
import base64, re, sys, tempfile
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path("C:/DATA/Projects/202606_E-PROFILE_calibration")
IMG = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
BOLD = re.compile(r"\*\*([^*]+)\*\*")
CODE = re.compile(r"`([^`]+)`")
MAXW = 6.3  # inches, page text width


def _add_runs(paragraph, text):
    """Add text to a paragraph, honouring **bold** and `code` inline spans."""
    # split on bold first; within non-bold, split on code
    pos = 0
    for m in BOLD.finditer(text):
        _add_code_runs(paragraph, text[pos:m.start()], bold=False)
        _add_code_runs(paragraph, m.group(1), bold=True)
        pos = m.end()
    _add_code_runs(paragraph, text[pos:], bold=False)


def _add_code_runs(paragraph, text, bold):
    pos = 0
    for m in CODE.finditer(text):
        r = paragraph.add_run(text[pos:m.start()]); r.bold = bold
        rc = paragraph.add_run(m.group(1)); rc.bold = bold; rc.font.name = "Consolas"; rc.font.size = Pt(9)
        pos = m.end()
    r = paragraph.add_run(text[pos:]); r.bold = bold


def _add_image(doc, src, base, report_dir):
    try:
        if src.startswith("data:"):
            b64 = src.split(",", 1)[1]
            data = base64.b64decode(b64)
            tmp = Path(tempfile.mkstemp(suffix=".png")[1]); tmp.write_bytes(data)
            path = tmp
        else:
            path = base / src
            if not path.is_file():
                path = report_dir / src
        if not Path(path).is_file():
            doc.add_paragraph(f"[missing image: {src[:60]}]")
            return
        doc.add_picture(str(path), width=Inches(MAXW))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as exc:  # noqa: BLE001
        doc.add_paragraph(f"[image error: {exc}]")


def _flush_table(doc, rows):
    # rows: list of list-of-cells; row[1] is the |---| separator -> skip it
    body = [r for i, r in enumerate(rows) if not (i == 1 and all(set(c) <= set("-: ") for c in r))]
    if not body:
        return
    ncol = max(len(r) for r in body)
    tbl = doc.add_table(rows=0, cols=ncol); tbl.style = "Light Grid Accent 1"
    for ri, r in enumerate(body):
        cells = tbl.add_row().cells
        for ci in range(ncol):
            txt = r[ci] if ci < len(r) else ""
            p = cells[ci].paragraphs[0]
            _add_runs(p, txt.strip())
            if ri == 0:
                for run in p.runs:
                    run.bold = True


def convert(md_path, out_path=None, base=BASE):
    md_path = Path(md_path)
    out_path = Path(out_path) if out_path else md_path.with_suffix(".docx")
    report_dir = md_path.parent
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"; doc.styles["Normal"].font.size = Pt(10.5)
    tbl_rows = []
    in_fence = False
    for raw in lines:
        line = raw.rstrip("\n")
        s = line.strip()
        # code fences: toggle; render contents as monospace
        if s.startswith("```"):
            if tbl_rows:
                _flush_table(doc, tbl_rows); tbl_rows = []
            in_fence = not in_fence
            continue
        if in_fence:
            p = doc.add_paragraph(); r = p.add_run(line)
            r.font.name = "Consolas"; r.font.size = Pt(8.5)
            continue
        # table rows: any pipe-bearing line that is not a heading/quote/list/image. Handles both
        # standard (| a | b |) and loose (a | b | c) GitHub pipe tables, and the |---| separator.
        if "|" in s and not s.startswith(("#", ">", "- ", "* ", "!")):
            tbl_rows.append([c for c in s.strip("|").split("|")])
            continue
        elif tbl_rows:
            _flush_table(doc, tbl_rows); tbl_rows = []
        # image (own line)
        m = IMG.match(s)
        if m:
            _add_image(doc, m.group(2), base, report_dir)
            cap = m.group(1).strip()
            if cap:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(cap); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            continue
        if not s:
            continue
        if s.startswith("#"):
            lvl = len(s) - len(s.lstrip("#"))
            doc.add_heading(s.lstrip("#").strip(), level=min(lvl, 4))
        elif s.startswith(">"):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            _add_runs(p, s.lstrip("> ").strip())
            for run in p.runs:
                run.italic = True
        elif s.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet"); _add_runs(p, s[2:].strip())
        elif re.match(r"^\d+\.\s", s):
            p = doc.add_paragraph(style="List Number"); _add_runs(p, re.sub(r"^\d+\.\s", "", s))
        elif set(s) <= set("-*_") and len(s) >= 3:
            continue  # horizontal rule
        else:
            p = doc.add_paragraph(); _add_runs(p, s)
    if tbl_rows:
        _flush_table(doc, tbl_rows)
    doc.save(str(out_path))
    print(f"wrote {out_path}  ({out_path.stat().st_size/1e3:.0f} kB)")
    return out_path


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2] if len(sys.argv) > 2 else None,
            Path(sys.argv[3]) if len(sys.argv) > 3 else BASE)
